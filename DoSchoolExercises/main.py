# Program created with the aim of making it easier to exercise and copy school materials at home and on the PC.
# Applications needed: Word and some IDE for Python.

# Importing library
from docx import Document

subject = str(input("\nDisciplina (Ex: Biologia): "))

print("","_"*18,"\n|"," "*3,"Conteúdo"," "*2,"|\n","_"*18,"\n| Exercícios  | 1 |\n| Matéria     | 2 |\n|","_"*15,"|")

content = int(input("\nTipo de conteúdo: "))

if content == 1:
    source = str(input("\nOs exercícios são do livro ou apostila? (s/n): "))
    if source == "s":
        pages = input("\nPáginas: ")
        numbers = input("Números: ")
        amount = int(input("Quantidade de Exercícios: "))
        filename = input("Nome do Arquivo: ")

        document = Document()
        document.add_heading(subject, 0)

        paragraph = document.add_paragraph() 
        paragraph.add_run("Conteúdo: ").bold = True
        paragraph.add_run("Exercícios").italic = True
        paragraph.add_run("\nPáginas: ").bold = True
        paragraph.add_run(f"{pages}").italic = True
        paragraph.add_run("\nNúmeros: ").bold = True
        paragraph.add_run(f"{numbers}").italic = True
        document.add_heading("", 0)

        for c in range(1,amount+1):
            paragraph2 = document.add_paragraph()
            paragraph2.add_run("\nPág.: ").bold = True
            paragraph2.add_run("\nNúm.: ").bold = True
            paragraph2.add_run("\nResposta: ").bold = True
            document.add_paragraph(" "*4)
            document.add_heading("", 0)
            
        document.save(f"{filename}.docx")

    elif source == "n":
        origin = input("Origem dos Exercícios (Ex: Site da Escola): ")
        amount = int(input("Quantidade de Exercícios: "))
        filename = input("Nome do Arquivo: ")

        document = Document()

        document.add_heading(subject, 0)

        paragraph = document.add_paragraph() 
        paragraph.add_run("Conteúdo: ").bold = True
        paragraph.add_run("Exercícios").italic = True
        paragraph.add_run("\nOrigem: ").bold = True
        paragraph.add_run(f"{origin}").italic = True
        paragraph.add_run("\nNúmero de Exercícios: ").bold = True
        paragraph.add_run(f"{amount}").italic = True
        document.add_heading("", 0)

        for c in range(1,amount+1):
            document.add_paragraph("\n[IMAGEM_EXERCICIO]", style='List Number')
            paragraph2 = document.add_paragraph()
            paragraph2.add_run("\nResposta: ").bold = True
            document.add_paragraph(" "*4)
            document.add_heading("", 0)

        document.save(f"{filename}.docx")


elif content == 2:
    title = input("Título: ")
    filename = input("Nome do Arquivo: ")

    document = Document()
    document.add_heading(f"{subject}", 0)
    document.add_paragraph(f"{title}", style="Intense Quote")
    document.add_paragraph(" ")
    document.save(f"{filename}.docx")
